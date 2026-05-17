"""
document_reader.py — Extraction de texte (version légère)
- PDF  : pdfplumber  (~2 Mo, pur Python, pas de dépendances C)
- DOCX : python-docx (~1 Mo)
- TXT  : intégré Python
"""

import io
import re
from typing import Tuple


def read_document(file_bytes: bytes, filename: str) -> Tuple[str, int]:
    """
    Extrait le texte d'un fichier selon son extension.
    Retourne: (texte_extrait, nombre_de_pages)
    """
    filename = filename.lower()
    if filename.endswith(".pdf"):
        return _read_pdf(file_bytes)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return _read_docx(file_bytes)
    elif filename.endswith(".txt") or filename.endswith(".rtf"):
        return _read_txt(file_bytes)
    else:
        raise ValueError(f"Format non supporté : {filename}")


def _read_pdf(file_bytes: bytes) -> Tuple[str, int]:
    """Extrait le texte d'un PDF avec pdfplumber (léger, pur Python)."""
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    pages_text.append(t.strip())
        full_text = "\n\n".join(pages_text)
        return _clean(full_text)[:50000], page_count
    except Exception as e:
        raise Exception(f"Erreur lecture PDF : {e}")


def _read_docx(file_bytes: bytes) -> Tuple[str, int]:
    """Extrait le texte d'un fichier Word (.docx)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_txt:
                    lines.append(row_txt)
        full_text = "\n".join(lines)
        full_text = _clean(full_text)
        page_count = max(1, len(full_text) // 3000)
        return full_text[:50000], page_count
    except Exception as e:
        raise Exception(f"Erreur lecture DOCX : {e}")


def _read_txt(file_bytes: bytes) -> Tuple[str, int]:
    """Lit un fichier texte brut."""
    try:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")
        text = _clean(text)
        return text[:50000], max(1, len(text) // 3000)
    except Exception as e:
        raise Exception(f"Erreur lecture TXT : {e}")


def _clean(text: str) -> str:
    """Nettoie le texte extrait."""
    text = text.replace("\x00", "")
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
